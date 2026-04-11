"""
Script que regenera el template Acta_Tecnica.docx con tags docxtpl correctos.
Usa python-docx para crear el .docx y luego modifica el XML para los loops.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import zipfile
import os
import shutil
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def crear_template():
    """Crea el template base con todos los tags docxtpl."""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    
    # TITULO
    titulo = doc.add_heading('ACTA T\xc9CNICA', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.size = Pt(16)
        run.font.name = 'Arial'
    
    # INFO GENERAL
    t_info = doc.add_table(rows=6, cols=2)
    t_info.style = 'Table Grid'
    campos = [
        ("N\xc2\xb0 Acta:", "{{numero_acta}}"),
        ("Fecha:", "{{fecha_larga}}"),
        ("Lugar:", "{{lugar}}"),
        ("Hora de inicio:", "{{hora_inicio}}"),
        ("Hora de finalizaci\xc3\xb3n:", "{{hora_fin}}"),
        ("Convocado por:", "{{convocante_nombre}}\n{{convocante_cargo}}"),
    ]
    for i, (label, valor) in enumerate(campos):
        p_l = t_info.cell(i, 0).paragraphs[0]
        p_l.paragraph_format.space_before = Pt(2)
        p_l.paragraph_format.space_after = Pt(2)
        r = p_l.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Arial'
        
        p_v = t_info.cell(i, 1).paragraphs[0]
        p_v.paragraph_format.space_before = Pt(2)
        p_v.paragraph_format.space_after = Pt(2)
        r = p_v.add_run(valor)
        r.font.size = Pt(10)
        r.font.name = 'Arial'
        t_info.cell(i, 0).width = Cm(4)
        t_info.cell(i, 1).width = Cm(12)
    
    doc.add_paragraph('')
    
    # PARTICIPANTES
    h_part = doc.add_heading('PARTICIPANTES', level=2)
    for run in h_part.runs:
        run.font.name = 'Arial'
    p_part = doc.add_paragraph("{{tabla_participantes}}")
    p_part.runs[0].font.size = Pt(10)
    p_part.runs[0].font.name = 'Arial'
    
    # DESARROLLO
    h_dev = doc.add_heading('DESARROLLO DE LA REUNI\xc3\x93N', level=2)
    for run in h_dev.runs:
        run.font.name = 'Arial'
    t_dev = doc.add_table(rows=3, cols=1)
    t_dev.style = 'Table Grid'
    for idx, txt in enumerate(["Puntos del orden del d\xc3\xada:", "{{aspectos_ia}}", "{{desarrollo_ia}}"]):
        p = t_dev.cell(idx, 0).paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt)
        r.font.size = Pt(10)
        r.font.name = 'Arial'
        if idx == 0:
            r.bold = True
    
    # COMPROMISOS
    h_comp = doc.add_heading('COMPROMISOS Y ACUERDOS', level=2)
    for run in h_comp.runs:
        run.font.name = 'Arial'
    p_comp = doc.add_paragraph("{{compromisos_ia}}")
    p_comp.runs[0].font.size = Pt(10)
    p_comp.runs[0].font.name = 'Arial'
    
    # EVIDENCIAS
    h_evid = doc.add_heading('EVIDENCIAS FOTOGR\xc3\x81FICAS', level=2)
    for run in h_evid.runs:
        run.font.name = 'Arial'
    p_evid = doc.add_paragraph("[Espacio para evidencias fotogr\xc3\xa1ficas]")
    p_evid.paragraph_format.space_before = Pt(20)
    p_evid.paragraph_format.space_after = Pt(20)
    
    # FIRMANTES - tabla con 2 rows: header + data
    h_firm = doc.add_heading('FIRMANTES', level=2)
    for run in h_firm.runs:
        run.font.name = 'Arial'
    
    t_firm = doc.add_table(rows=2, cols=2)
    t_firm.style = 'Table Grid'
    
    # Header
    for j, txt in enumerate(["NOMBRE Y CARGO", "FIRMA"]):
        p = t_firm.cell(0, j).paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Arial'
    
    # Fila de datos (se modificar\xc3\xa1 luego con tags de loop)
    p_nombre = t_firm.cell(1, 0).paragraphs[0]
    p_nombre.paragraph_format.space_before = Pt(2)
    p_nombre.paragraph_format.space_after = Pt(2)
    r = p_nombre.add_run("{{f.nombre}}\n{{f.cargo}}")
    r.font.size = Pt(10)
    r.font.name = 'Arial'
    
    t_firm.cell(1, 1).paragraphs[0].add_run("")
    t_firm.cell(1, 0).width = Cm(10)
    t_firm.cell(1, 1).width = Cm(6)
    
    doc.add_paragraph('')
    
    # ELABORADO POR
    p_elab = doc.add_paragraph()
    p_elab.paragraph_format.space_before = Pt(12)
    r_elab_label = p_elab.add_run("Elaborado por: ")
    r_elab_label.bold = True
    r_elab_label.font.size = Pt(10)
    r_elab_label.font.name = 'Arial'
    r_elab_val = p_elab.add_run("{{elaborado_titulo}} {{elaborado_nombre}}")
    r_elab_val.font.size = Pt(10)
    r_elab_val.font.name = 'Arial'
    
    # Guardar temporal
    out_path = os.path.join(os.path.dirname(__file__), '..', 'resources', 'Acta_Tecnica.docx')
    bak_path = out_path + '.bak'
    
    # Backup del original
    if os.path.exists(out_path):
        shutil.copy2(out_path, bak_path)
    
    doc.save(out_path)
    
    # ═══════════════════════════════════════════════════
    # Ahora modificamos el XML para agregar los tags de loop
    # ═══════════════════════════════════════════════════
    
    # Extraer document.xml
    with zipfile.ZipFile(out_path, 'r') as z:
        xml_content = z.read('word/document.xml')
    
    # Parsear
    tree = etree.fromstring(xml_content)
    
    # Encontrar la \xc3\xbaltima tabla (firmantes)
    tables = tree.findall('.//w:tbl', {'w': W_NS})
    t_firm_xml = tables[-1]
    rows_xml = t_firm_xml.findall('w:tr', {'w': W_NS})
    data_row_xml = rows_xml[1]  # La fila de datos (indice 1)
    
    # Crear fila de apertura del loop
    open_row = etree.SubElement(t_firm_xml, f'{{{W_NS}}}tr')
    open_cell = etree.SubElement(open_row, f'{{{W_NS}}}tc')
    open_p = etree.SubElement(open_cell, f'{{{W_NS}}}p')
    open_r = etree.SubElement(open_p, f'{{{W_NS}}}r')
    open_t = etree.SubElement(open_r, f'{{{W_NS}}}t')
    open_t.text = '{%tr for f in firmantes %}'
    
    # Modificar la fila de datos: limpiar celda 0 y poner tags
    cells = data_row_xml.findall('w:tc', {'w': W_NS})
    cell0 = cells[0]
    # Limpiar parrafos existentes
    for p in cell0.findall('w:p', {'w': W_NS}):
        cell0.remove(p)
    
    # Parrafo con {{f.nombre}}
    p1 = etree.SubElement(cell0, f'{{{W_NS}}}p')
    r1 = etree.SubElement(p1, f'{{{W_NS}}}r')
    t1 = etree.SubElement(r1, f'{{{W_NS}}}t')
    t1.text = '{{f.nombre}}'
    
    # Parrafo con {{f.cargo}}
    p2 = etree.SubElement(cell0, f'{{{W_NS}}}p')
    r2 = etree.SubElement(p2, f'{{{W_NS}}}r')
    t2 = etree.SubElement(r2, f'{{{W_NS}}}t')
    t2.text = '{{f.cargo}}'
    
    # Crear fila de cierre del loop
    close_row = etree.SubElement(t_firm_xml, f'{{{W_NS}}}tr')
    close_cell = etree.SubElement(close_row, f'{{{W_NS}}}tc')
    close_p = etree.SubElement(close_cell, f'{{{W_NS}}}p')
    close_r = etree.SubElement(close_p, f'{{{W_NS}}}r')
    close_t = etree.SubElement(close_r, f'{{{W_NS}}}t')
    close_t.text = '{%tr endfor %}'
    
    # Convertir de vuelta a XML
    xml_str = etree.tostring(tree, encoding='UTF-8', xml_declaration=True, pretty_print=True)
    
    # Reescribir el docx
    with zipfile.ZipFile(out_path, 'a') as z:
        z.writestr('word/document.xml', xml_str)
    
    print(f"\u2705 Template generado exitosamente en: {out_path}")
    print("")
    print("Tags configurados:")
    print("  - {{numero_acta}}, {{fecha_larga}}, {{lugar}}, {{hora_inicio}}, {{hora_fin}}")
    print("  - {{convocante_nombre}}, {{convocante_cargo}}")
    print("  - {{tabla_participantes}}")
    print("  - {{aspectos_ia}}, {{desarrollo_ia}}, {{compromisos_ia}}")
    print("  - {{elaborado_titulo}}, {{elaborado_nombre}}")
    print("  - {%tr for f in firmantes %} ... {{f.nombre}}, {{f.cargo}} ... {%tr endfor %}")

if __name__ == '__main__':
    crear_template()

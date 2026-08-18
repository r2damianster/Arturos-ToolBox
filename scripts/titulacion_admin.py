"""
Administración del catálogo de Titulación — Evaluación de Pares Lectores.
Permite dar de alta una modalidad/rúbrica nueva sin tocar la lógica del wizard
ni del generador de documentos: ambos leen el catálogo de forma genérica
(logic/titulacion_db.py), así que solo hace falta insertar filas nuevas.

Uso:
    python scripts/titulacion_admin.py inspeccionar "ruta/al/archivo.docx"
        Imprime la estructura de tablas del .docx (filas/columnas/texto) para
        ayudar a armar el 'schema' del JSON de configuración.

    python scripts/titulacion_admin.py registrar "ruta/al/config.json"
        Inserta o actualiza (por slug) la modalidad + rúbrica descritas en el
        JSON. Ver scripts/ejemplo_modalidad.json para el formato esperado.
"""
import sys
import os
import json
import argparse

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))

from docx import Document
from logic.titulacion_db import init_titulacion_db, upsert_modalidad, upsert_rubrica

ESCALAS_VALIDAS = {"peso_si_no", "si_no", "niveles_4", "niveles_especial"}


def inspeccionar(ruta_docx):
    doc = Document(ruta_docx)
    print(f"Archivo: {ruta_docx}")
    print(f"Tablas encontradas: {len(doc.tables)}\n")
    for tabla_idx, tabla in enumerate(doc.tables):
        print(f"--- Tabla {tabla_idx}: {len(tabla.rows)} filas x {len(tabla.columns)} columnas ---")
        for fila_idx, fila in enumerate(tabla.rows):
            celdas = [c.text.strip().replace("\n", " | ") for c in fila.cells]
            print(f"  fila {fila_idx}: {celdas}")
        print()
    print(
        "Con esta estructura arma el campo 'schema' del JSON de configuración: cada tabla "
        "necesita 'nombre', 'escala' (" + " | ".join(sorted(ESCALAS_VALIDAS)) + ") y su lista "
        "de 'criterios' (con 'texto' y 'peso' cuando la escala lo requiera).\n"
        "Ver scripts/ejemplo_modalidad.json para un ejemplo completo."
    )


def _validar_schema(schema):
    if "tablas" not in schema:
        raise ValueError("El schema debe tener una clave 'tablas'.")
    for tabla in schema["tablas"]:
        if tabla.get("escala") not in ESCALAS_VALIDAS:
            raise ValueError(f"Escala inválida en tabla '{tabla.get('nombre')}': {tabla.get('escala')}")
        if not tabla.get("criterios"):
            raise ValueError(f"La tabla '{tabla.get('nombre')}' no tiene criterios.")


def registrar(ruta_config):
    with open(ruta_config, encoding="utf-8") as f:
        config = json.load(f)

    modalidad_cfg = config["modalidad"]
    rubrica_cfg = config["rubrica"]
    _validar_schema(rubrica_cfg["schema"])

    init_titulacion_db()  # asegura que las tablas ya existan (no reseedea si ya hay datos)

    modalidad_id = upsert_modalidad(
        modalidad_cfg["slug"],
        modalidad_cfg["nombre"],
        modalidad_cfg.get("requiere_subtipo", False),
    )
    rubrica_id = upsert_rubrica(
        modalidad_id,
        rubrica_cfg["slug"],
        rubrica_cfg.get("subtipo"),
        rubrica_cfg["plantilla_docx"],
        rubrica_cfg["schema"],
    )

    print(f"Modalidad '{modalidad_cfg['slug']}' -> id {modalidad_id}")
    print(f"Rúbrica '{rubrica_cfg['slug']}' -> id {rubrica_id}")
    print("Listo. El wizard y el generador de documentos ya la reconocen (son data-driven).")


def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    sub = parser.add_subparsers(dest="comando", required=True)

    p1 = sub.add_parser("inspeccionar", help="Imprime la estructura de tablas de un .docx")
    p1.add_argument("archivo_docx")

    p2 = sub.add_parser("registrar", help="Registra/actualiza modalidad+rúbrica desde un JSON de config")
    p2.add_argument("archivo_config")

    args = parser.parse_args()
    if args.comando == "inspeccionar":
        inspeccionar(args.archivo_docx)
    elif args.comando == "registrar":
        registrar(args.archivo_config)


if __name__ == "__main__":
    main()

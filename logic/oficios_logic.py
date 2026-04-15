"""
Lógica para generación de Oficios con IA.
Genera oficios dirigidos a cualquier persona de la base de datos,
con contenido enriquecido por IA y tono personalizable.
"""
import os
import io
import datetime
import re
from docx import Document
from logic.db import get_all_docentes, get_all_carreras


class OficioLogic:
    """Genera oficios en formato .docx usando la plantilla HOJA_CARRERA_PINE.docx."""

    def __init__(self):
        self.base_dir = os.path.dirname(os.path.abspath(__file__))
        self.resources_path = os.path.normpath(os.path.join(self.base_dir, "..", "resources"))
        self.template_name = "HOJA_CARRERA_PINE.docx"

    # ── Tonos disponibles ──────────────────────────────────────────
    TONOS = {
        "formal": {
            "system": "Eres un asistente de redacción administrativa universitaria. Tu tono es extremadamente formal, respetuoso y protocolar.",
            "instruction": "Usa lenguaje institucional elevado. Incluye fórmulas de cortesía académica. Mantén distancia protocolar.",
        },
        "cordial": {
            "system": "Eres un asistente de redacción administrativa universitaria. Tu tono es formal pero cercano y amable.",
            "instruction": "Usa lenguaje respetuoso pero cálido. Incluye expresiones de colaboración y trabajo en equipo.",
        },
        "directo": {
            "system": "Eres un asistente de redacción administrativa universitaria. Tu tono es directo y conciso, sin rodeos.",
            "instruction": "Ve al grano. Usa oraciones cortas. Elimina preámbulos innecesarios. Mantén la formalidad básica.",
        },
        "urgente": {
            "system": "Eres un asistente de redacción administrativa universitaria. Tu tono transmite urgencia e importancia.",
            "instruction": "Usa lenguaje que destaca la prioridad y los plazos. Mantén formalidad pero con sentido de inmediatez.",
        },
    }

    # ── Utilidades de texto ────────────────────────────────────────
    def formatear_fecha_larga(self, fecha_str):
        """Convierte 2026-04-15 en '15 de abril de 2026'."""
        try:
            if not fecha_str or " de " in str(fecha_str):
                return fecha_str
            fecha_obj = datetime.datetime.strptime(str(fecha_str), '%Y-%m-%d')
            meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                     "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
            return f"{fecha_obj.day} de {meses[fecha_obj.month - 1]} de {fecha_obj.year}"
        except Exception:
            return fecha_str

    def reemplazar_en_documento(self, doc, reemplazos):
        """Reemplaza etiquetas {{LLAVE}} ignorando mayúsculas/minúsculas en todo el documento."""
        limpios = {}
        for k, v in reemplazos.items():
            key_name = str(k).replace('{', '').replace('}', '').strip()
            val = str(v) if v is not None else ""
            limpios[key_name.upper()] = val

        def realizar_cambio(texto):
            nuevo_texto = texto
            for match in re.findall(r'\{\{(.*?)\}\}', nuevo_texto):
                tag_interna = match.strip().upper()
                if tag_interna in limpios:
                    nuevo_texto = nuevo_texto.replace("{{" + match + "}}", limpios[tag_interna])
            return nuevo_texto

        # Párrafos
        for p in doc.paragraphs:
            if '{{' in p.text:
                p.text = realizar_cambio(p.text)

        # Tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if '{{' in p.text:
                            p.text = realizar_cambio(p.text)

    # ── Generación del oficio ──────────────────────────────────────
    def generar_docx(self, datos):
        """
        Genera el .docx del oficio.

        Args:
            datos: dict con las siguientes claves:
                - num_oficio: número del oficio
                - fecha_emision: fecha en formato YYYY-MM-DD
                - ciudad: ciudad de emisión
                - destinatario_nombre: nombre completo del destinatario
                - destinatario_cargo: cargo del destinatario
                - destinatario_carrera: carrera/dependencia del destinatario
                - asunto: asunto del oficio
                - cuerpo: cuerpo principal del oficio
                - firmante_titulo: título del firmante
                - firmante_nombre: nombre del firmante
                - firmante_cargo: cargo del firmante
                - iniciales: iniciales del elaborador
                - tono: tono usado por IA (formal, cordial, directo, urgente)

        Returns:
            BytesIO con el documento .docx generado.
        """
        template_path = os.path.join(self.resources_path, self.template_name)
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")

        doc = Document(template_path)

        # Normalizar datos
        datos_normalizados = {}
        for k, v in datos.items():
            key_limpia = k.replace('{', '').replace('}', '')
            if 'fecha' in key_limpia.lower():
                v = self.formatear_fecha_larga(v)
            datos_normalizados[k] = v

        # Reemplazar en el documento
        self.reemplazar_en_documento(doc, datos_normalizados)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def obtener_destinatarios(self, carrera=None):
        """
        Obtiene lista de destinatarios disponibles.
        Si carrera es None, devuelve todos.
        """
        if carrera:
            from logic.db import get_docentes_by_carrera
            return get_docentes_by_carrera(carrera)
        return get_all_docentes()

    def obtener_carreras(self):
        """Obtiene lista de carreras/dependencias disponibles."""
        return get_all_carreras()

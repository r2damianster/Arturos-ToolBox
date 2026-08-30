"""
Generador de los 2 documentos de salida del par lector (Titulación):
1) Informe de Criterios Observados (plantilla única, agnóstica de modalidad)
2) Rúbrica de calificación (data-driven desde schema_json — sirve para TEFL y Artículo)

Se generan desde cero con python-docx: los .docx originales en Proyecto_Titulacion/
tienen celdas fusionadas de forma inconsistente (verificado con python-docx) y no
tienen membrete/logo que preservar, así que replicar su estructura a mano es más
confiable que editar el archivo original celda por celda.
"""
import io
import datetime
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _fecha_larga(fecha=None):
    fecha = fecha or datetime.date.today()
    meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
             "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    return f"{fecha.day} de {meses[fecha.month - 1]} de {fecha.year}"


def _indicadores_por_clave(indicadores):
    return {(i['tabla_idx'], i['criterio_idx']): i for i in indicadores}


# ── Documento 1: Informe de Criterios Observados ──────────────────────────

def generar_informe_docx(evaluacion):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)

    doc.add_paragraph(f"MEMORANDUM No. {evaluacion.get('numero_memo') or '—'}").runs[0].bold = True
    doc.add_paragraph("PARA: Miembros de Comisión Académica")
    doc.add_paragraph("ASUNTO: Criterios observados en el trabajo de integración curricular y/o examen complexivo")
    doc.add_paragraph(f"FECHA: Manta, {_fecha_larga()}")
    doc.add_paragraph()

    tutor = evaluacion.get('tutor') or '__________________'
    doc.add_paragraph(
        "En cumplimiento a lo que dispone el Reglamento de Régimen Académico Interno sobre el "
        "proceso de titulación, una vez que se ha revisado el trabajo de integración curricular "
        f"y/o examen complexivo para el cual fue designado/a en dirigir {tutor}, de acuerdo al "
        "siguiente detalle:"
    )

    tabla1 = doc.add_table(rows=2, cols=4)
    tabla1.style = 'Table Grid'
    encabezados1 = ["Facultad y/o Extensión", "Carrera", "Opción de Titulación", "Título"]
    valores1 = [
        evaluacion.get('facultad') or '',
        evaluacion.get('carrera') or '',
        evaluacion.get('opcion_titulacion') or '',
        evaluacion.get('titulo_trabajo') or '',
    ]
    for celda, texto in zip(tabla1.rows[0].cells, encabezados1):
        celda.paragraphs[0].add_run(texto).bold = True
    for celda, texto in zip(tabla1.rows[1].cells, valores1):
        celda.text = texto

    doc.add_paragraph()
    opcion = evaluacion.get('opcion_titulacion') or 'Trabajo de Integración Curricular o Examen Complexivo'
    doc.add_paragraph(
        "Luego de haber realizado el análisis en cada uno de los componentes que forman parte del "
        "trabajo escrito y en concordancia con una de las competencias otorgadas al tribunal de "
        "titulación —que consiste en que, luego de la revisión del trabajo, éste deberá emitir un "
        "informe con las observaciones de forma y contenido sobre el documento presentado, el mismo "
        "que deberá ser conocido por el tutor/a quien direccionará al estudiante para que realice las "
        f"correcciones necesarias— se detallan los criterios del {opcion} que fueron observados:"
    )

    observaciones = evaluacion.get('observaciones') or []
    formales = [o for o in observaciones if o['seccion'] == 'formal']
    fondo = [o for o in observaciones if o['seccion'] == 'fondo']

    tabla2 = doc.add_table(rows=1, cols=2)
    tabla2.style = 'Table Grid'
    tabla2.rows[0].cells[0].paragraphs[0].add_run("Componentes").bold = True
    tabla2.rows[0].cells[1].paragraphs[0].add_run("Observaciones").bold = True

    def _agregar_seccion(titulo, filas):
        fila = tabla2.add_row()
        fila.cells[0].paragraphs[0].add_run(titulo).bold = True
        for obs in filas:
            fila = tabla2.add_row()
            fila.cells[0].text = obs['componente']
            fila.cells[1].text = obs.get('observacion') or ''

    _agregar_seccion("Aspectos formales:", formales)
    _agregar_seccion("Aspectos de fondo:", fondo)

    doc.add_paragraph()
    doc.add_paragraph("Particular que se informa para los fines consiguientes.")
    doc.add_paragraph()
    doc.add_paragraph("Atentamente,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(evaluacion.get('evaluador_nombre') or '__________________')
    doc.add_paragraph("Miembro del Tribunal Calificador")
    doc.add_paragraph(f"Correo Electrónico Institucional: {evaluacion.get('evaluador_correo') or ''}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ── Documento 2: Rúbrica (data-driven desde plantilla y schema_json) ──────────────────

def obtener_ruta_plantilla(plantilla_nombre):
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    candidatos = [
        os.path.join(base_dir, plantilla_nombre),
        os.path.join(base_dir, "resources", plantilla_nombre),
        os.path.join(base_dir, "Proyecto_Titulacion", "TEFL", plantilla_nombre),
        os.path.join(base_dir, "Proyecto_Titulacion", "Artifuclo", plantilla_nombre),
    ]
    for c in candidatos:
        if os.path.exists(c):
            return c
    raise FileNotFoundError(f"No se encontró la plantilla de rúbrica: {plantilla_nombre}")


def generar_rubrica_docx(evaluacion):
    import os
    rubrica = evaluacion.get('rubrica')
    if not rubrica:
        raise ValueError("La evaluación no tiene una rúbrica asignada.")
    schema = rubrica['schema']
    indicadores = _indicadores_por_clave(evaluacion.get('indicadores') or [])
    puntaje_total = evaluacion.get('puntaje_total', 0)

    # 1. Cargar la plantilla original
    plantilla_nombre = rubrica['plantilla_docx']
    ruta_plantilla = obtener_ruta_plantilla(plantilla_nombre)
    doc = Document(ruta_plantilla)

    # 2. Reemplazar encabezados en los párrafos de texto
    estudiante = evaluacion.get('estudiante') or ''
    evaluador = evaluacion.get('evaluador_nombre') or ''
    fecha_txt = _fecha_larga()
    tutor = evaluacion.get('tutor') or ''
    tema = evaluacion.get('titulo_trabajo') or ''
    escala_total = schema.get('escala_total', 10)
    puntaje_formateado = f"{puntaje_total:.2f}" if isinstance(puntaje_total, (int, float)) else str(puntaje_total)

    for p in doc.paragraphs:
        p_text = p.text
        if not p_text.strip():
            continue

        if "Estudiante:" in p_text or "ESTUDIANTE:" in p_text:
            for run in p.runs:
                if "_" in run.text:
                    run.text = " " + estudiante
                    break
        elif "Evaluador:" in p_text or "MIEMBRO DEL TRIBUNAL:" in p_text or "Tribunal:" in p_text:
            for run in p.runs:
                if "_" in run.text:
                    run.text = " " + evaluador
                    break
        elif "Tutor:" in p_text or "TUTOR:" in p_text:
            for run in p.runs:
                if "_" in run.text:
                    run.text = " " + tutor
                    break
        elif "Tema:" in p_text or "TEMA:" in p_text or "TITULO DEL ARTICULO" in p_text:
            for run in p.runs:
                if "_" in run.text:
                    run.text = " " + tema
                    break
        elif "Fecha:" in p_text or "FECHA:" in p_text:
            has_calificacion = "Calificaci" in p_text or "Total" in p_text
            if has_calificacion:
                p.text = f"Fecha: {fecha_txt}                 Calificación Total: {puntaje_formateado} / {escala_total}"
                p.runs[0].bold = True
            else:
                for run in p.runs:
                    if "_" in run.text:
                        run.text = " " + fecha_txt
                        break
        elif "FECHA DE ENTREGA:" in p_text:
            for run in p.runs:
                if "_" in run.text:
                    run.text = " " + fecha_txt
                    break

    # 3. Rellenar las tablas de la plantilla
    if doc.tables:
        table = doc.tables[0]
        tabla_schema = schema['tablas'][0]
        escala = tabla_schema['escala']
        header_rows = tabla_schema.get('header_rows', 1)
        tabla_idx = 0

        for criterio_idx, criterio in enumerate(tabla_schema['criterios']):
            indicador = indicadores.get((tabla_idx, criterio_idx), {})
            row_idx = criterio_idx + header_rows

            if row_idx >= len(table.rows):
                break

            row = table.rows[row_idx]
            respuesta = indicador.get('respuesta')
            calificacion = indicador.get('calificacion')
            comentario = indicador.get('comentario') or ''

            if escala in ('peso_si_no', 'si_no'):
                num_cols = len(row.cells)
                if num_cols >= 6:
                    # YES (Col 3)
                    cell_yes = row.cells[3]
                    cell_yes.text = 'X' if respuesta == 'YES' else ''
                    cell_yes.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # NO (Col 4)
                    cell_no = row.cells[4]
                    cell_no.text = 'X' if respuesta == 'NO' else ''
                    cell_no.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Puntaje (Col 5)
                    cell_score = row.cells[5]
                    if calificacion is not None:
                        cell_score.text = f"{calificacion:.2f}"
                    else:
                        cell_score.text = ""
                    cell_score.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif escala in ('niveles_4', 'niveles_especial'):
                num_cols = len(row.cells)
                if num_cols >= 8:
                    for c in range(2, 6):
                        row.cells[c].text = ""

                    col_x = None
                    if respuesta == '0':
                        col_x = 2
                    elif respuesta == '35':
                        col_x = 3
                    elif respuesta in ('70', '90'):
                        col_x = 4
                    elif respuesta == '100':
                        col_x = 5

                    if col_x is not None:
                        row.cells[col_x].text = 'X'
                        row.cells[col_x].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Calif. (Col 6)
                    cell_score = row.cells[6]
                    if calificacion is not None:
                        cell_score.text = f"{calificacion:.2f}"
                    else:
                        cell_score.text = ""
                    cell_score.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Comentarios (Col 7)
                    row.cells[7].text = comentario

        # 4. Rellenar fila de Total
        for row in table.rows:
            if len(row.cells) > 0 and row.cells[0].text.strip().upper() in ('TOTAL', 'TOTALES'):
                num_cols = len(row.cells)
                if escala in ('peso_si_no', 'si_no') and num_cols >= 6:
                    cell_total = row.cells[5]
                    cell_total.text = puntaje_formateado
                    cell_total.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif escala in ('niveles_4', 'niveles_especial') and num_cols >= 7:
                    cell_total = row.cells[6]
                    cell_total.text = puntaje_formateado
                    cell_total.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

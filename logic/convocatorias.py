import os
import io
import pandas as pd
import unicodedata
import datetime
import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from logic.db import get_all_docentes

class ConvocatoriaLogic:
    def __init__(self):
        self.base_dir = os.path.dirname(os.path.abspath(__file__))
        self.resources_path = os.path.normpath(os.path.join(self.base_dir, "..", "resources"))

    def normalizar(self, texto):
        """Elimina acentos y prepara texto para ordenamiento."""
        if not texto: return ""
        texto = str(texto).lower()
        texto = unicodedata.normalize('NFD', texto)
        return ''.join(c for c in texto if unicodedata.category(c) != 'Mn')

    def formatear_fecha_reunion(self, fecha_str):
        """Convierte 2026-02-20 en '20 de febrero de 2026'."""
        try:
            if not fecha_str or " de " in str(fecha_str): 
                return fecha_str
            
            fecha_obj = datetime.datetime.strptime(str(fecha_str), '%Y-%m-%d')
            meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", 
                     "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
            return f"{fecha_obj.day} de {meses[fecha_obj.month - 1]} de {fecha_obj.year}"
        except:
            return fecha_str

    def reemplazar_en_documento(self, doc, reemplazos):
        """Reemplaza etiquetas {{LLAVE}} ignorando mayúsculas/minúsculas."""
        limpios = {}
        for k, v in reemplazos.items():
            key_name = str(k).replace('{', '').replace('}', '').strip()
            val = str(v) if v is not None else ""
            limpios[key_name.upper()] = val

        def realizar_cambio(texto):
            nuevo_texto = texto
            # Busca cualquier cosa dentro de {{ }}
            for match in re.findall(r'\{\{(.*?)\}\}', nuevo_texto):
                tag_interna = match.strip().upper()
                if tag_interna in limpios:
                    nuevo_texto = nuevo_texto.replace("{{" + match + "}}", limpios[tag_interna])
            return nuevo_texto

        # Procesar párrafos
        for p in doc.paragraphs:
            if '{{' in p.text:
                p.text = realizar_cambio(p.text)
        
        # Procesar tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if '{{' in p.text:
                            p.text = realizar_cambio(p.text)

    def insertar_estudiantes_en_todas_las_tablas(self, doc, lista_estudiantes):
        """Llena todas las tablas que contengan el marcador {{nombre}}."""
        for table in doc.tables:
            marcador_encontrado = False
            fila_plantilla = None
            
            for row in table.rows:
                texto_celda = row.cells[2].text.lower() if len(row.cells) > 2 else ""
                if "{{nombre}}" in texto_celda:
                    marcador_encontrado = True
                    fila_plantilla = row
                    break
            
            if marcador_encontrado:
                parent = fila_plantilla._tr.getparent()
                parent.remove(fila_plantilla._tr)
                
                for i, estudiante in enumerate(lista_estudiantes, start=1):
                    new_row = table.add_row()
                    new_row.cells[0].text = str(i)
                    new_row.cells[1].text = "Estudiante"
                    new_row.cells[2].text = estudiante
                    new_row.cells[3].text = ""

    def _identificar_columnas_nombres(self, df):
        """Busca índices de columnas que parezcan nombres y apellidos."""
        idx_nombre = 0
        idx_apellido = 1
        
        cols = [str(c).lower() for c in df.columns]
        
        # Prioridad 1: Buscar "apellido" y "nombre"
        for i, c in enumerate(cols):
            if 'apellido' in c:
                idx_apellido = i
            elif 'nombre' in c:
                idx_nombre = i
                
        # Si son iguales (no encontró uno), intentar ser más específico
        if idx_nombre == idx_apellido:
            # Si solo encontró "nombre" pero no "apellido", tal vez es col 0 y 1
            idx_nombre = 0
            idx_apellido = 1
            
        return idx_nombre, idx_apellido

    def procesar_excel_estudiantes(self, lista_archivos):
        """Une múltiples excels, limpia y ordena nombres."""
        nombres_consolidados = []
        print(f"DEBUG: Procesando {len(lista_archivos)} archivos de estudiantes.")
        
        for excel_file in lista_archivos:
            try:
                # 1. Asegurar puntero al inicio
                excel_file.seek(0)
                
                # 2. Leer según extensión
                ext = os.path.splitext(excel_file.filename)[1].lower()
                engine = 'odf' if ext == '.ods' else None
                
                df = pd.read_excel(excel_file, engine=engine)
                
                if df.empty:
                    print(f"ADVERTENCIA: El archivo {excel_file.filename} está vacío.")
                    continue
                
                # 3. Identificar columnas
                idx_n, idx_a = self._identificar_columnas_nombres(df)
                
                count = 0
                for _, row in df.iterrows():
                    if len(row) <= max(idx_n, idx_a): continue
                    
                    nombre = str(row.iloc[idx_n]).strip()
                    apellido = str(row.iloc[idx_a]).strip()
                    
                    if nombre and apellido and "nan" not in (nombre.lower() + apellido.lower()):
                        # Evitar que los encabezados entren como nombres si se detectaron mal
                        if nombre.lower() == "nombre" or apellido.lower() == "apellido":
                            continue
                            
                        nombre_completo = f"{apellido} {nombre}".upper()
                        nombres_consolidados.append(nombre_completo)
                        count += 1
                
                print(f"DEBUG: Archivo '{excel_file.filename}' -> {count} estudiantes encontrados.")
                
            except Exception as e:
                print(f"Error procesando archivo '{getattr(excel_file, 'filename', 'S/N')}': {e}")
                continue
        
        finales = sorted(list(set(nombres_consolidados)), key=self.normalizar)
        print(f"DEBUG: Total consolidado: {len(finales)} estudiantes.")
        return finales

    def generar_docx(self, tipo, datos, excel_files=None, docentes_manuales=None):
        """Punto de entrada principal."""
        filename = "Convocatoria_Docentes.docx" if tipo == "docente" else "Convocatoria_Estudiantes.docx"
        path = os.path.join(self.resources_path, filename)

        if not os.path.exists(path):
            raise FileNotFoundError(f"No se encontró la plantilla en: {path}")

        doc = Document(path)

        # 1. Normalizar las llaves de 'datos' para que la fecha se procese bien
        datos_normalizados = {}
        for k, v in datos.items():
            key_limpia = k.replace('{','').replace('}','').lower()
            if 'fecha_reunion' in key_limpia:
                v = self.formatear_fecha_reunion(v)
            datos_normalizados[k] = v

        # 2. Si es docente, agregar datos dinámicos desde la BD o manuales
        if tipo == "docente":
            if docentes_manuales:
                # Modo manual: normalizar claves
                docentes_norm = []
                for d in docentes_manuales:
                    docentes_norm.append({
                        'titulo_grado': d.get('titulo', d.get('titulo_grado', 'Lic.')),
                        'nombre': d.get('nombre', ''),
                        'post_grado': d.get('postgrado', d.get('post_grado', '')),
                        'cargo': d.get('cargo', 'Docente'),
                        'carrera': d.get('carrera', ''),
                        'es_director': d.get('es_director', False),
                    })
                datos_docentes = {'docentes': docentes_norm}
            else:
                # Modo carrera: cargar desde la BD
                datos_docentes = self._generar_datos_docentes()
            datos_normalizados.update(datos_docentes)
            # Expandir filas dinámicas en las tablas
            self._expandir_tabla_lista(doc, datos_docentes['docentes'])
            self._expandir_tabla_firmas(doc, datos_docentes['docentes'])

        # 3. Reemplazar texto (incluyendo placeholders de filas expandidas)
        self.reemplazar_en_documento(doc, datos_normalizados)

        # 4. Si hay estudiantes, procesar tablas
        if tipo == "estudiante" and excel_files:
            estudiantes = self.procesar_excel_estudiantes(excel_files)
            if estudiantes:
                self.insertar_estudiantes_en_todas_las_tablas(doc, estudiantes)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _generar_datos_docentes(self):
        """Genera el diccionario de reemplazos para los placeholders del template de docentes."""
        docentes = get_all_docentes()
        datos = {'docentes': docentes}
        # Placeholders estáticos (no numerados) que se reemplazan en cada fila
        return datos

    def _expandir_tabla_lista(self, doc, docentes):
        """Expande la tabla de asistencia: duplica la fila plantilla por cada docente."""
        for table in doc.tables:
            # Detectar si es la tabla de lista (tiene encabezado con "NOMBRES")
            header = [c.text.strip() for c in table.rows[0].cells]
            if 'NOMBRES' not in header:
                continue

            # Buscar la fila plantilla (contiene {{NOMBRE_DOCENTE}})
            template_row = None
            template_idx = None
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue  # skip header
                cell_text = row.cells[2].text  # columna NOMBRES
                if '{{NOMBRE_DOCENTE}}' in cell_text:
                    template_row = row
                    template_idx = i
                    break

            if template_row is None:
                return

            # Remover la fila plantilla
            table._tbl.remove(template_row._tr)

            # Agregar una fila por cada docente
            for idx, d in enumerate(docentes, start=1):
                nombre_completo = f"{d['titulo_grado']} {d['nombre']}, {d['post_grado']}"
                new_row = table.add_row()
                new_row.cells[0].text = str(idx)
                new_row.cells[1].text = d['cargo']
                new_row.cells[2].text = nombre_completo
                new_row.cells[3].text = ''

            return  # solo una tabla de lista

    def _expandir_tabla_firmas(self, doc, docentes):
        """Expande la tabla de firmas: duplica la fila plantilla según cantidad de docentes."""
        for table in doc.tables:
            # Detectar si es la tabla de firmas (tiene "He sido convocado:")
            header = [c.text.strip() for c in table.rows[0].cells]
            if 'He sido convocado:' not in header:
                continue

            # Buscar la fila plantilla (contiene {{FIRMA_NOMBRE}})
            template_row = None
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue
                for cell in row.cells:
                    if '{{FIRMA_NOMBRE}}' in cell.text:
                        template_row = row
                        break
                if template_row:
                    break

            if template_row is None:
                return

            # Remover la fila plantilla
            table._tbl.remove(template_row._tr)

            # Calcular cuántas filas necesitamos (2 firmas por fila)
            cols = len(table.columns)
            total = len(docentes)
            num_rows = (total + cols - 1) // cols  # ceiling division

            for r in range(num_rows):
                new_row = table.add_row()
                for c in range(cols):
                    doc_idx = r * cols + c
                    if doc_idx < total:
                        d = docentes[doc_idx]
                        nombre_completo = f"{d['titulo_grado']} {d['nombre']}, {d['post_grado']}"
                        # Rebuild cell content from template
                        for p in new_row.cells[c].paragraphs:
                            p.clear()
                        p0 = new_row.cells[c].paragraphs[0]
                        r1 = p0.add_run('___________________________')
                        br = r1._r.makeelement(qn('w:br'), {})
                        r1._r.append(br)
                        p0.add_run(nombre_completo)
                        br2 = p0.add_run('')
                        br2._r.makeelement(qn('w:br'), {})
                        r2_el = r1._r.makeelement(qn('w:br'), {})
                        p0._p.append(r2_el)
                        p0.add_run(d['cargo'])
                    else:
                        # Empty cell
                        for p in new_row.cells[c].paragraphs:
                            p.clear()

            return  # solo una tabla de firmas
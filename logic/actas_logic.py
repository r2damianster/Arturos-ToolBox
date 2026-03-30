import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches
from groq import Groq

GROQ_API_KEY = os.environ.get('GROQ_API_KEY', '')

class ActaTecnicaLogic:
    def __init__(self):
        self.model_id = "llama-3.3-70b-versatile"
        try:
            self.client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None
        except Exception as e:
            print(f"Error crítico de inicialización: {e}")
            self.client = None

    def generar_texto_ia(self, tipo, notas_usuario):
        if not notas_usuario or len(notas_usuario.strip()) < 2:
            return "No se registraron detalles adicionales para esta sección."
        if not self.client:
            return f"[IA no configurada] {notas_usuario}"

        if tipo == "aspectos":
            seccion_desc = "Puntos del orden del día"
            instruccion = (
                "TAREA: Organiza los puntos del orden del día.\n"
                "REGLAS: No expandas el texto innecesariamente. Solo corrige ortografía y formaliza levemente el vocabulario.\n"
                "FORMATO: Presenta cada punto precedido por una viñeta (•). No redactes párrafos largos."
            )
            max_tokens, temperature = 200, 0.2
        elif tipo == "compromisos":
            seccion_desc = "Decisiones y compromisos finales"
            instruccion = (
                "TAREA: Redacta los acuerdos y compromisos institucionales.\n"
                "REGLAS: Sé directo y breve. Mantén la esencia de la nota original sin añadir relleno.\n"
                "FORMATO: Presenta cada compromiso precedido por una viñeta (•). Corrige coherencia y ortografía."
            )
            max_tokens, temperature = 200, 0.2
        else:
            seccion_desc = "Desarrollo y deliberaciones de la reunión"
            instruccion = (
                "TAREA: Convierte las notas en un relato académico fluido y profesional.\n"
                "REGLAS: Aquí SÍ puedes expandirte. Conecta los puntos del orden del día con conectores lógicos.\n"
                "FORMATO: Redacta en párrafos narrativos. Usa tono solemne (ej: 'asimismo', 'se procedió a')."
            )
            max_tokens, temperature = 600, 0.5

        prompt = (
            f"Actúa como un Secretario Académico universitario de alto nivel.\n\n"
            f"SECCIÓN: {seccion_desc}.\n"
            f"{instruccion}\n\n"
            f"NOTAS DEL USUARIO:\n{notas_usuario}"
        )
        try:
            completion = self.client.chat.completions.create(
                model=self.model_id,
                messages=[
                    {"role": "system", "content": "Eres un redactor experto que diferencia entre listas breves y desarrollo narrativo."},
                    {"role": "user",   "content": prompt}
                ],
                temperature=temperature,
                max_tokens=max_tokens
            )
            return completion.choices[0].message.content.strip().replace('**', '').replace('*', '•')
        except Exception as e:
            return f"Error en la generación de texto: {str(e)}. Notas: {notas_usuario}"

    def formatear_fecha(self, fecha_str):
        try:
            dt = datetime.strptime(fecha_str, '%Y-%m-%d')
            meses = ["enero","febrero","marzo","abril","mayo","junio",
                     "julio","agosto","septiembre","octubre","noviembre","diciembre"]
            return f"{dt.day} de {meses[dt.month - 1]} de {dt.year}"
        except:
            return fecha_str

    def _reconstruir_tabla_firmas(self, tabla, convocante_nombre, convocante_cargo, participantes):
        """Reconstruye la tabla de firmas con convocante + participantes (2 columnas)."""
        # Row 0 ya existe: col0 = convocante, col1 = primer participante o vacío
        # Llenamos row 0
        def set_cell(cell, nombre, cargo):
            for p in cell.paragraphs:
                p.clear()
            cell.paragraphs[0].add_run(nombre)
            cell.add_paragraph(cargo)
            cell.add_paragraph('')  # espacio para firma

        set_cell(tabla.cell(0, 0), convocante_nombre, convocante_cargo)

        # Distribuir participantes: primer slot es col1 de row0, luego filas nuevas de 2
        slots = [('row0_col1',)] + [('new',)] * len(participantes)

        if participantes:
            set_cell(tabla.cell(0, 1), participantes[0][0], participantes[0][1])
        else:
            tabla.cell(0, 1).paragraphs[0].clear()

        # Agregar filas nuevas de a 2 participantes
        restantes = participantes[1:]
        for i in range(0, len(restantes), 2):
            new_row = tabla.add_row()
            set_cell(new_row.cells[0], restantes[i][0], restantes[i][1])
            if i + 1 < len(restantes):
                set_cell(new_row.cells[1], restantes[i+1][0], restantes[i+1][1])

    def _insertar_fotos(self, tabla_evidencias, fotos):
        """Inserta imágenes en la celda de evidencias fotográficas."""
        if not fotos:
            return
        # Tabla 3 tiene [0]=header desarrollo, [1]=desarrollo_ia, [2]=Evidencias
        cell = tabla_evidencias.cell(2, 0)
        cell.paragraphs[0].clear()
        for foto in fotos:
            try:
                img_stream = io.BytesIO(foto.read())
                para = cell.add_paragraph()
                run  = para.add_run()
                run.add_picture(img_stream, width=Inches(2.5))
            except Exception as e:
                cell.add_paragraph(f"[Error al insertar imagen: {e}]")

    def crear_docx(self, form_data, fotos=None):
        convocante_raw = form_data.get('convocante', '')
        partes_conv    = convocante_raw.split(',')
        nombre_conv    = partes_conv[0].strip() if len(partes_conv) > 0 else "Convocante"
        cargo_conv     = partes_conv[1].strip() if len(partes_conv) > 1 else ""

        # Participantes con cargo
        titulos   = form_data.getlist('p_titulo[]')
        nombres   = form_data.getlist('p_nombre[]')
        apellidos = form_data.getlist('p_apellido[]')
        cargos    = form_data.getlist('p_cargo[]')
        participantes = []
        for t, n, a, c in zip(titulos, nombres, apellidos, cargos):
            if n:
                nombre_completo = f"{t} {n} {a}".strip()
                participantes.append((nombre_completo, c or "Docente"))

        # Elaborado por
        elab_titulo = form_data.get('elaborado_titulo', '').strip()
        elab_nombre = form_data.get('elaborado_nombre', '').strip()

        # Textos IA
        asp = self.generar_texto_ia("aspectos",    form_data.get('notas_aspectos', ''))
        des = self.generar_texto_ia("desarrollo",  form_data.get('notas_reunion', ''))
        com = self.generar_texto_ia("compromisos", form_data.get('notas_compromisos', ''))

        fecha_larga = self.formatear_fecha(form_data.get('fecha_reunion', ''))

        logic_dir     = os.path.dirname(os.path.abspath(__file__))
        root_dir      = os.path.dirname(logic_dir)
        template_path = os.path.join(root_dir, 'resources', 'Acta_Tecnica.docx')

        doc = Document(template_path)

        # Reemplazos de texto simples
        reemplazos = {
            "{{numero_acta}}":          form_data.get('num_acta', 'S/N'),
            "{{fecha_larga}}":          fecha_larga,
            "{{lugar}}":               form_data.get('lugar_reunion', 'Instalaciones Institucionales'),
            "{{hora_inicio}}":         form_data.get('hora_inicio', '--:--'),
            "{{hora_fin}}":            form_data.get('hora_fin', '--:--'),
            "{{convocante_nombre}}":   nombre_conv,
            "{{convocante_cargo}}":    cargo_conv,
            "{{tabla_participantes}}": "\n".join(f"{p[0]} — {p[1]}" for p in participantes),
            "{{aspectos_ia}}":         asp,
            "{{desarrollo_ia}}":       des,
            "{{compromisos_ia}}":      com,
            "{{elaborado_titulo}}":    elab_titulo,
            "{{elaborado_nombre}}":    elab_nombre,
            "{{firma_0_1}}":           participantes[0][0] if participantes else "",
        }

        for p in doc.paragraphs:
            for k, v in reemplazos.items():
                if k in p.text:
                    p.text = p.text.replace(k, str(v))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for k, v in reemplazos.items():
                            if k in paragraph.text:
                                paragraph.text = paragraph.text.replace(k, str(v))

        # Reconstruir tabla de firmas (tabla índice 6)
        self._reconstruir_tabla_firmas(
            doc.tables[6], nombre_conv, cargo_conv, participantes
        )

        # Insertar fotos de evidencia (tabla índice 3)
        if fotos:
            self._insertar_fotos(doc.tables[3], fotos)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
